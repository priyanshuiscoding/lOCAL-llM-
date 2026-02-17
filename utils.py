import os
import fitz  # PyMuPDF
import pandas as pd
import docx
from PIL import Image
import pytesseract
import psycopg2
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.llms import Ollama

try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter

try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.docstore.document import Document

UPLOAD_FOLDER = "uploads"
INDEX_FOLDER = "faiss_index"

# Save uploaded file to disk
def save_file(uploaded_file):
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
    with open(path, "wb") as f:
        f.write(uploaded_file.read())
    return path

# Load uploaded file
def load_file(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _load_pdf(path)
    elif ext == ".txt":
        return _load_txt(path)
    elif ext == ".csv":
        return _load_csv(path)
    elif ext == ".docx":
        return _load_docx(path)
    elif ext in [".jpg", ".jpeg", ".png"]:
        return _load_image(path)
    else:
        raise ValueError("Unsupported file type")

# Text extractors
def _load_pdf(path):
    doc = fitz.open(path)
    texts = [page.get_text() for page in doc]
    return [Document(page_content=t) for t in texts if t.strip()]

def _load_txt(path):
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    return [Document(page_content=text)]

def _load_csv(path):
    df = pd.read_csv(path)
    text = df.to_string()
    return [Document(page_content=text)]

def _load_docx(path):
    doc = docx.Document(path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return [Document(page_content=text)]

def _load_image(path):
    image = Image.open(path)
    text = pytesseract.image_to_string(image)
    return [Document(page_content=text)]

# ✅ Load from PostgreSQL (now uses mytestdb)
def load_from_postgres(
    host="localhost",
    dbname="mytestdb",
    user="postgres",
    password="5386",
    table="employees"
):
    try:
        conn = psycopg2.connect(
            host=host,
            dbname=dbname,
            user=user,
            password=password
        )
        cursor = conn.cursor()
        cursor.execute(f"SELECT * FROM {table}")
        rows = cursor.fetchall()
        colnames = [desc[0] for desc in cursor.description]

        documents = []
        for row in rows:
            row_text = ", ".join(f"{col}: {val}" for col, val in zip(colnames, row))
            documents.append(Document(page_content=row_text))

        cursor.close()
        conn.close()
        return documents
    except Exception as e:
        print("❌ Error loading from PostgreSQL:")
        print(e)
        return []

# Split docs into chunks
def chunk_documents(documents):
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=150
    )
    return text_splitter.split_documents(documents)

# Save chunks into FAISS
def save_to_vectordb(chunks):
    embeddings = HuggingFaceEmbeddings()
    vectordb = FAISS.from_documents(chunks, embeddings)
    vectordb.save_local(INDEX_FOLDER)

# Query FAISS + local LLM (Ollama)
def query_vectordb(query):
    embeddings = HuggingFaceEmbeddings()
    vectordb = FAISS.load_local(INDEX_FOLDER, embeddings, allow_dangerous_deserialization=True)
    docs = vectordb.similarity_search(query)

    if not docs:
        return "❌ No relevant information found in your uploaded files."

    llm = Ollama(model="mistral")  # Or llama2, gemma, etc.
    context = "\n\n".join(doc.page_content for doc in docs[:5])
    prompt = (
        "Use the context below to answer the question.\n\n"
        f"Context:\n{context}\n\n"
        f"Question:\n{query}\n\n"
        "Answer:"
    )
    result = llm.invoke(prompt)
    return result

# ✅ Optional: Load and embed PostgreSQL data at startup
if __name__ == "__main__":
    docs = load_from_postgres()
    chunks = chunk_documents(docs)
    save_to_vectordb(chunks)
    print("✅ PostgreSQL data loaded and embedded into vector store.")
